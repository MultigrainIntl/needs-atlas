#!/usr/bin/env python3
"""
generate_assessment.py — Generate a county Community Needs Assessment in Zufall's
house style.

County + statewide HEADLINE figures come from config/real_counties.json (Census
QuickFacts — the accurate published county/state values). When the ACS pipeline
has run (data/tracts_need.csv present), the report is ENRICHED with a real
census-tract concentration analysis — the thing county averages hide: how many
neighborhoods carry above-average need, how much of the county's population lives
in them, and where the sharpest need sits.

County + tract figures come from Census ACS; the county food-insecurity figure
comes from Feeding America's Map the Meal Gap (config/food_insecurity.json).
Indicators that require external sources not yet wired in (United Way ALICE,
provider-to-population ratios, chronic-disease prevalence) are listed honestly
as pending — never estimated or fabricated.

    python scripts/generate_assessment.py Middlesex
"""
import csv, json, re, sys, datetime
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
R = json.loads((ROOT / "config" / "real_counties.json").read_text())
STATE = R["state"]
BY_NAME = {c["county"]: c for c in R["counties"]}
TRACTS = ROOT / "data" / "tracts_need.csv"
ROLLUP = ROOT / "data" / "counties.json"
TODAY = datetime.date.today().strftime("%B %d, %Y")

_FI_PATH = ROOT / "config" / "food_insecurity.json"
FI = json.loads(_FI_PATH.read_text()) if _FI_PATH.exists() else {"counties": {}}
FIC = FI.get("counties", {})

_PL_PATH = ROOT / "config" / "places.json"
PL = json.loads(_PL_PATH.read_text()) if _PL_PATH.exists() else {}
PLC = PL.get("counties", {})
PLNJ = PL.get("nj", {})
PLLAB = PL.get("measures", {})

_PR_PATH = ROOT / "config" / "providers.json"
PR = json.loads(_PR_PATH.read_text()) if _PR_PATH.exists() else {}
PRC = PR.get("counties", {})
PRNJ = PR.get("nj", {})

_AL_PATH = ROOT / "config" / "alice.json"
AL = json.loads(_AL_PATH.read_text()) if _AL_PATH.exists() else {}
ALC = AL.get("counties", {})
ALNJ = AL.get("nj", {})

_MT_PATH = ROOT / "config" / "maternal.json"
MT = json.loads(_MT_PATH.read_text()) if _MT_PATH.exists() else {}
MTC = MT.get("counties", {})
MTNJ = MT.get("nj", {})

CITES = {
    "qf": ("U.S. Census Bureau QuickFacts. {county} County, New Jersey and New Jersey. "
           "2020-2024 American Community Survey 5-Year Estimates and 2025 Population "
           "Estimates. Accessed " + TODAY + "."),
    "acs": ("U.S. Census Bureau, American Community Survey 2020-2024 5-Year Estimates, "
            "census-tract tables B27001 (health insurance), C17002 (ratio of income to "
            "poverty), B17001 (poverty), C16002 (household language), B22003 (SNAP receipt), "
            "retrieved via the Needs Atlas data pipeline and aggregated to the tract level. "
            "Accessed " + TODAY + "."),
    "feeding": ("Feeding America. Map the Meal Gap — {year} county food-insecurity estimates. "
                "Retrieved from map.feedingamerica.org. Accessed " + TODAY + ".").format(
                    year=FI.get("data_year", "")),
    "places": ("Centers for Disease Control and Prevention. PLACES: Local Data for Better Health, "
               "County Data, " + str(PL.get("release", "")) + " release — model-based small-area "
               "estimates (crude prevalence) from the Behavioral Risk Factor Surveillance System. "
               "Retrieved from cdc.gov/places. Accessed " + TODAY + "."),
    "chr": ("County Health Rankings & Roadmaps, " + str(PR.get("release", "")) + " release. "
            "University of Wisconsin Population Health Institute. Retrieved from countyhealthrankings.org. "
            "Accessed " + TODAY + "."),
    "alice": ("United Way — United For ALICE. New Jersey ALICE Report (" + str(AL.get("data_year", "")) +
              " data). Retrieved from unitedforalice.org. Accessed " + TODAY + "."),
}

# External sources not yet wired in — listed honestly as pending, never fabricated.
PENDING_ALL = [
    ("njshad", "Prenatal care and infant mortality, including racial disparities",
     "New Jersey State Health Assessment Data (NJSHAD)"),
]

CITE_RE = re.compile(r"\[\[cite:([a-z_]+)\]\]")


def num(r, k):
    try:
        v = float(r[k]); return v
    except (ValueError, KeyError, TypeError):
        return None


def rel(cv, sv, hi="above", lo="below", eq="comparable to"):
    return eq if abs(cv - sv) < max(0.4, 0.03 * sv) else (hi if cv > sv else lo)


def oxford(items):
    items = list(items)
    if not items:
        return ""
    if len(items) == 1:
        return items[0]
    if len(items) == 2:
        return f"{items[0]} and {items[1]}"
    return ", ".join(items[:-1]) + ", and " + items[-1]


# ---------------------------------------------------------------------------
# Tract analysis (real ACS data)
# ---------------------------------------------------------------------------
def load_tracts():
    if not TRACTS.exists():
        return None
    return [r for r in csv.DictReader(TRACTS.open())]


def wmean(rows, k):
    n = d = 0.0
    for r in rows:
        v, p = num(r, k), num(r, "total_pop")
        if v is not None and p:
            n += v * p; d += p
    return (n / d) if d else None


def is_artifact(r):
    """Flag likely group-quarters / non-residential tracts so they never headline a
    'highest-need neighborhood'. Signals: 0% uninsured with ~100% below 200% FPL
    (dorms, facilities), missing insurance universe, or a very small population."""
    u = num(r, "uninsured_pct"); f200 = num(r, "under_200_fpl_pct"); pop = num(r, "total_pop")
    if pop is not None and pop < 1200:
        return True
    if u is not None and u == 0.0 and f200 is not None and f200 >= 90.0:
        return True
    if u is None or f200 is None:
        return True
    return False


def concentration(county, tracts, ref):
    """Return dict of real concentration stats for one county, or None."""
    rs = [r for r in tracts if r["county_name"] == county]
    if not rs:
        return None
    pop = sum(num(r, "total_pop") or 0 for r in rs)
    su, sf = ref["uninsured_pct"], ref["under_200_fpl_pct"]
    n_unins = sum(1 for r in rs if (num(r, "uninsured_pct") or 0) > su)
    n_f200 = sum(1 for r in rs if (num(r, "under_200_fpl_pct") or 0) > sf)
    pop_hi = sum((num(r, "total_pop") or 0) for r in rs
                 if (num(r, "under_200_fpl_pct") or 0) >= 25)
    scores = [num(r, "need_score") for r in rs if num(r, "need_score") is not None]
    real = [r for r in rs if not is_artifact(r)]
    spot = sorted(real, key=lambda r: num(r, "need_score") or 0, reverse=True)[:5]
    return {
        "n": len(rs), "pop": int(pop),
        "n_unins": n_unins, "n_f200": n_f200,
        "pop_hi": int(pop_hi), "pop_hi_pct": round(100 * pop_hi / pop) if pop else 0,
        "smin": round(min(scores)) if scores else None,
        "smax": round(max(scores)) if scores else None,
        "spot": spot,
    }


def county_rank():
    """Map county -> (rank, total) by ACS need_index, 1 = highest need."""
    if not ROLLUP.exists():
        return {}
    cs = json.loads(ROLLUP.read_text())["counties"]
    order = sorted(cs, key=lambda c: c.get("need_index", 0), reverse=True)
    return {c["county"]: (i + 1, len(cs)) for i, c in enumerate(order)}


ORD = {1: "highest", 2: "second-highest", 3: "third-highest", 4: "fourth-highest",
       5: "fifth-highest", 6: "sixth-highest", 7: "lowest"}


# ---------------------------------------------------------------------------
# Narrative
# ---------------------------------------------------------------------------
def build(c, tracts, ref, ranks):
    s = STATE
    conc = concentration(c["county"], tracts, ref) if tracts else None
    has_tracts = conc is not None
    aln = (ALC.get(c["county"]) or {}).get("pct_below_alice")
    alice_clause = (f"Poverty alone understates the strain: {aln}% of {c['county']} County households fall below "
                    f"the ALICE threshold — unable to afford basic necessities despite, in most cases, being "
                    f"employed (New Jersey: {ALNJ.get('pct_below_alice')}%).[[cite:alice]] ") if aln is not None else ""
    secs = [
        ("Poverty, income, and the case for concentrated need",
         f"The communities Zufall Health serves within {c['county']} County, New Jersey "
         f"include areas of concentrated poverty and constrained access to care. "
         f"{c['county']} County's median household income is ${c['median_hh_income']:,}, "
         f"{rel(c['median_hh_income'], s['median_hh_income'], 'above', 'below', 'near')} the "
         f"statewide median of ${s['median_hh_income']:,}, and {c['poverty_pct']}% of "
         f"residents live in poverty, {rel(c['poverty_pct'], s['poverty_pct'])} the New Jersey "
         f"rate of {s['poverty_pct']}%.[[cite:qf]] {alice_clause}County averages, however, flatten the "
         f"neighborhoods where need concentrates — which a census-tract view brings into focus."),
    ]

    if has_tracts:
        rk = ranks.get(c["county"])
        rank_line = (f" Across the seven counties Zufall serves, {c['county']} carries the "
                     f"{ORD.get(rk[0], str(rk[0]) + 'th')} composite need." if rk else "")
        spot = conc["spot"]
        spot_txt = ""
        if spot:
            t = spot[0]
            lep = num(t, "lep_pct")
            lep_clause = (f", and {lep:.0f}% of households are limited-English-speaking"
                          if lep is not None else "")
            short = f"{int(t['GEOID'][-6:]) / 100:.2f}".rstrip("0").rstrip(".")
            town = (t.get("municipality") or "").strip()
            place_clause = (f"{town} (Census Tract {short}, GEOID {t['GEOID']})"
                            if town else f"Census Tract {short} (GEOID {t['GEOID']})")
            spot_txt = (f" The sharpest need sits in {place_clause}: there, "
                        f"{num(t,'uninsured_pct'):.0f}% of residents are uninsured and "
                        f"{num(t,'under_200_fpl_pct'):.0f}% live below 200% of the federal "
                        f"poverty level{lep_clause} — a concentration that county-level averages "
                        f"hide entirely.")
        if conc["pop_hi"] > 0:
            conc_line = (f"{conc['pop_hi_pct']}% of the county's residents — roughly "
                         f"{conc['pop_hi']:,} people — live in tracts where at least one in four "
                         f"residents falls below 200% of the federal poverty level.")
        else:
            conc_line = ("No census tract in the county reaches the threshold of a quarter or more "
                         "residents below 200% of the poverty level — need here is lower and more "
                         "evenly distributed than in Zufall's higher-need counties.")
        secs.append((
            "Where the need concentrates: a census-tract view",
            f"{c['county']} County contains {conc['n']} census tracts. Measured against the "
            f"seven-county service-area average, {conc['n_unins']} of those tracts carry an "
            f"above-average uninsured rate and {conc['n_f200']} carry an above-average share of "
            f"residents below 200% of the federal poverty level.[[cite:acs]] "
            f"{conc_line}{spot_txt}{rank_line} This is precisely the sub-county geography a "
            f"grant application must document, and it is where Zufall's sites and outreach can be "
            f"targeted for greatest effect."))

    prov = PRC.get(c["county"])
    _rat = lambda x: (x or "").replace(":1", "")
    prov_txt = ""
    if prov:
        prov_txt = (f" Provider supply reinforces the gap: {c['county']} County has roughly one primary-care "
                    f"physician per {_rat(prov['pcp_ratio'])} residents, one dentist per {_rat(prov['dent_ratio'])}, "
                    f"and one mental-health provider per {_rat(prov['mh_ratio'])} — against statewide figures of "
                    f"{_rat(PRNJ['pcp_ratio'])}, {_rat(PRNJ['dent_ratio'])}, and {_rat(PRNJ['mh_ratio'])} "
                    f"respectively.[[cite:chr]]")
    secs.append((
        "Barriers to care: insurance and providers",
        f"Access to coverage is a clear barrier. {c['uninsured_under65_pct']}% of {c['county']} "
        f"County residents under age 65 are uninsured, "
        f"{rel(c['uninsured_under65_pct'], s['uninsured_under65_pct'])} the statewide rate of "
        f"{s['uninsured_under65_pct']}%.[[cite:qf]]" + prov_txt))

    h = PLC.get(c["county"])
    if h and h.get("DIABETES") is not None:
        elevated = []
        for m in ("DIABETES", "OBESITY", "BPHIGH", "HIGHCHOL", "CHD", "STROKE",
                  "CASTHMA", "COPD", "CANCER", "ARTHRITIS", "DEPRESSION"):
            cv, nv = h.get(m), PLNJ.get(m)
            if cv is not None and nv is not None and cv > nv:
                elevated.append(PLLAB.get(m, m).lower())
        if elevated:
            elev_clause = (f"{c['county']} exceeds the statewide rate on {len(elevated)} of the "
                           f"eleven tracked chronic conditions, including {oxford(elevated[:3])}. ")
        else:
            elev_clause = ("The county tracks at or below the statewide average across the tracked "
                           "chronic conditions. ")
        mlbw = (MTC.get(c["county"]) or {}).get("low_birthweight_pct")
        lbw_txt = (f" Birth outcomes reflect the same pressures: {mlbw}% of live births in {c['county']} "
                   f"County are low birthweight (New Jersey: {MTNJ.get('low_birthweight_pct')}%).[[cite:chr]]") if mlbw is not None else ""
        secs.append((
            "Chronic disease and community health burden",
            f"Model-based CDC estimates put adult diabetes in {c['county']} County at "
            f"{h['DIABETES']}% (New Jersey: {PLNJ['DIABETES']}%), obesity at {h['OBESITY']}% "
            f"(NJ: {PLNJ['OBESITY']}%), and high blood pressure at {h['BPHIGH']}% "
            f"(NJ: {PLNJ['BPHIGH']}%).[[cite:places]] {elev_clause}"
            f"{h['GHLTH']}% of adults report fair or poor overall health (NJ: {PLNJ['GHLTH']}%) and "
            f"{h['MHLTH']}% report frequent mental distress (NJ: {PLNJ['MHLTH']}%). On health-related "
            f"social needs, {h['LACKTRPT']}% of adults report a lack of reliable transportation and "
            f"{h['HOUSINSECU']}% report housing insecurity (NJ: {PLNJ['LACKTRPT']}% and "
            f"{PLNJ['HOUSINSECU']}%). These conditions define the day-to-day clinical demand a "
            f"community health center must meet — and, paired with the coverage and language "
            f"barriers documented above, the case for sustained primary, dental, and "
            f"behavioral-health capacity." + lbw_txt))

    fi = FIC.get(c["county"])
    if fi and fi.get("food_insecurity_pct") is not None:
        ppl = fi.get("food_insecure_people")
        ppl_clause = f" — about {ppl:,} people —" if ppl else ""
        child = fi.get("child_food_insecurity_pct")
        child_clause = (f", and {child}% of the county's children face food insecurity"
                        if child is not None else "")
        secs.append((
            "Food insecurity",
            f"An estimated {fi['food_insecurity_pct']}% of {c['county']} County residents"
            f"{ppl_clause} live in food-insecure households{child_clause}.[[cite:feeding]] Food "
            f"insecurity compounds the barriers documented here: it tracks closely with low income "
            f"and uninsurance and drives demand for the wraparound support Zufall's sites provide."))

    secs.append((
        "Demographics and language access",
        f"{c['county']} County is home to {c['population']:,} residents. {c['pct_hispanic']}% "
        f"identify as Hispanic or Latino and {c['pct_black']}% as Black or African American, "
        f"compared with {s['pct_hispanic']}% and {s['pct_black']}% statewide.[[cite:qf]] Language "
        f"access is a substantial need: {c['language_other_pct']}% of residents age 5 and over "
        f"speak a language other than English at home, "
        f"{rel(c['language_other_pct'], s['language_other_pct'])} the New Jersey figure of "
        f"{s['language_other_pct']}%.[[cite:qf]] Adults age 65 and over make up {c['pct_65_plus']}% "
        f"of the population (New Jersey: {s['pct_65_plus']}%).[[cite:qf]]"))

    return secs, PENDING_ALL, conc


def render_runs(secs):
    order = []
    def keynum(k):
        if k not in order: order.append(k)
        return str(order.index(k) + 1)
    out = []
    for title, text in secs:
        runs, pos = [], 0
        for m in CITE_RE.finditer(text):
            runs.append(("t", text[pos:m.start()])); runs.append(("c", keynum(m.group(1)))); pos = m.end()
        runs.append(("t", text[pos:]))
        out.append((title, runs))
    return out, order


def spotlight_rows(conc):
    """Return header + top-tract rows for the appendix table."""
    hdr = ["Municipality", "Census tract (GEOID)", "Uninsured", "Below 200% FPL", "Lim.-Eng. HH", "Need score"]
    rows = []
    for t in (conc["spot"] if conc else []):
        def g(k, suf="%"):
            v = num(t, k); return f"{v:.0f}{suf}" if v is not None else "—"
        short = f"{int(t['GEOID'][-6:]) / 100:.2f}".rstrip("0").rstrip(".")
        rows.append([(t.get("municipality") or "—").strip(),
                     f"{short}  ({t['GEOID']})", g("uninsured_pct"), g("under_200_fpl_pct"),
                     g("lep_pct"), g("need_score", "")])
    return hdr, rows


# ---------------------------------------------------------------------------
# Renderers
# ---------------------------------------------------------------------------
def to_docx(c, secs, order, pending, conc, path):
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    doc = Document()
    sec = doc.sections[0]; sec.page_width, sec.page_height = Inches(8.5), Inches(11)
    sec.left_margin = sec.right_margin = Inches(1)
    doc.styles["Normal"].font.name = "Calibri"; doc.styles["Normal"].font.size = Pt(11)
    r = doc.add_paragraph().add_run(f"{c['county']} County Needs Assessment — {datetime.date.today():%B %Y}")
    r.bold = True; r.font.size = Pt(16)
    note = ("County and statewide figures below are published U.S. Census Bureau values; the "
            "census-tract concentration analysis is drawn from ACS 5-year tract tables; food insecurity "
            "is from Feeding America's Map the Meal Gap; chronic-disease and community-health figures are "
            "from the CDC's PLACES project; ALICE household data is from United For ALICE; and "
            "provider-to-population ratios are from County Health Rankings. Remaining indicators (see end) "
            "are listed as pending — not estimated.")
    b = doc.add_paragraph().add_run(note)
    b.italic = True; b.font.size = Pt(8.5); b.font.color.rgb = RGBColor(0x1E, 0x6B, 0x57)
    for title, runs in secs:
        doc.add_paragraph().add_run(title).bold = True
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(10)
        for kind, txt in runs:
            run = p.add_run(txt)
            if kind == "c": run.font.superscript = True

    # Real tract appendix table
    if conc and conc["spot"]:
        doc.add_paragraph().add_run("Highest-need census tracts (ACS 5-year)").bold = True
        hdr, rows = spotlight_rows(conc)
        tbl = doc.add_table(rows=1, cols=len(hdr)); tbl.style = "Light Grid Accent 1"
        for j, h in enumerate(hdr):
            cell = tbl.rows[0].cells[j]; cell.text = ""
            rr = cell.paragraphs[0].add_run(h); rr.bold = True; rr.font.size = Pt(9)
        for row in rows:
            cells = tbl.add_row().cells
            for j, val in enumerate(row):
                cells[j].text = ""; rr = cells[j].paragraphs[0].add_run(val); rr.font.size = Pt(9)
        cap = doc.add_paragraph().add_run(
            "Group-quarters tracts (e.g., dormitories or institutional populations) are excluded "
            "so the ranking reflects residential neighborhoods.")
        cap.italic = True; cap.font.size = Pt(8)

    doc.add_paragraph().add_run("Indicators pending source connection").bold = True
    for _, label, src in pending:
        pp = doc.add_paragraph(style="List Bullet")
        pp.add_run(label + " — ").font.size = Pt(10)
        em = pp.add_run(src); em.italic = True; em.font.size = Pt(10)

    doc.add_paragraph().add_run("References").bold = True
    for i, k in enumerate(order, 1):
        rp = doc.add_paragraph(f"{i}. " + CITES[k].format(county=c["county"]))
        for run in rp.runs: run.font.size = Pt(8.5)
    doc.save(path)


def to_md(c, secs, order, pending, conc):
    out = [f"# {c['county']} County Needs Assessment — {datetime.date.today():%B %Y}", "",
           "_County/state figures are published Census values; tract concentration from ACS 5-year "
           "tract tables; food insecurity from Feeding America; chronic disease/community health from "
           "CDC PLACES; ALICE from United For ALICE; provider ratios from County Health Rankings. "
           "Remaining indicators (listed below) pending — not estimated._\n"]
    for title, runs in secs:
        out.append(f"## {title}")
        out.append("".join(t if k == "t" else f"[{t}]" for k, t in runs) + "\n")
    if conc and conc["spot"]:
        out.append("## Highest-need census tracts (ACS 5-year)")
        hdr, rows = spotlight_rows(conc)
        out.append("| " + " | ".join(hdr) + " |")
        out.append("| " + " | ".join(["---"] * len(hdr)) + " |")
        for row in rows:
            out.append("| " + " | ".join(row) + " |")
        out.append("\n_Group-quarters tracts (dormitories/institutional) excluded so the ranking "
                   "reflects residential neighborhoods._\n")
    out.append("## Indicators pending source connection")
    out += [f"- {label} — *{src}*" for _, label, src in pending]
    out.append("\n## References")
    out += [f"{i}. " + CITES[k].format(county=c["county"]) for i, k in enumerate(order, 1)]
    return "\n".join(out)


def generate(name):
    if name not in BY_NAME:
        raise SystemExit(f"Unknown county '{name}'. Options: {', '.join(BY_NAME)}")
    c = BY_NAME[name]
    tracts = load_tracts()
    ref = None
    if tracts:
        ref = {"uninsured_pct": round(wmean(tracts, "uninsured_pct"), 1),
               "under_200_fpl_pct": round(wmean(tracts, "under_200_fpl_pct"), 1)}
    ranks = county_rank()
    secs, pending, conc = build(c, tracts, ref, ranks)
    runs, order = render_runs(secs)
    outdir = ROOT / "data"; outdir.mkdir(exist_ok=True)
    stem = name.lower() + "_needs_assessment"
    (outdir / f"{stem}.md").write_text(to_md(c, runs, order, pending, conc))
    to_docx(c, runs, order, pending, conc, outdir / f"{stem}.docx")
    return stem


if __name__ == "__main__":
    name = sys.argv[1] if len(sys.argv) > 1 else "Middlesex"
    stem = generate(name)
    print(f"Wrote {stem}.docx/.md for {name} County"
          + (" (with real tract concentration analysis)" if TRACTS.exists()
             else " (county-level; run pipeline for tract detail)"))
